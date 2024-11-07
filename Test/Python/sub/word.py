from docx import Document
from pdfdocument.document import PDFDocument


class Docx:
    def __init__(self, file_name='example.docx'):
        self.file_name = file_name
        self.doc = Document()

    #生成docx文件
    def create_doc(self):
        # 添加标题
        self.doc.add_heading('Document Title', level=1)

        # 添加段落
        self.doc.add_paragraph('This is the first paragraph of the document.')

        # 添加带样式的段落
        paragraph = self.doc.add_paragraph('This is a ')
        paragraph.add_run('bold and italic ').bold = True
        paragraph.add_run('text. ').italic = True

        # 添加编号列表
        self.doc.add_paragraph('List Item 1', style='List Number')
        self.doc.add_paragraph('List Item 2', style='List Number')

        # 添加项目符号列表
        self.doc.add_paragraph('Bullet Point 1', style='List Bullet')
        self.doc.add_paragraph('Bullet Point 2', style='List Bullet')

        # 保存文档
        self.doc.save(self.file_name)

        return self.file_name

    # 将docx转换为pdf
    def docx_to_pdf(self, input_file, output_file):
        doc = Document(input_file)
        pdf = PDFDocument(output_file)
        pdf.init_report()

        for paragraph in doc.paragraphs:
            pdf.p(paragraph.text)

        for table in doc.tables:
            with pdf.table(style="border: 1px solid black; border-collapse: collapse;"):
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    pdf.tr(*row_data)

        pdf.generate()

        return output_file
